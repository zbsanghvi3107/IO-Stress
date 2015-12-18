####################################################
#                  Revision: 1.0                   #
#              Updated on: 11/06/2015              #
####################################################

####################################################
#                                                  #
#   This file writes/append data to the Template   #
#   Word File.                                     #
#                                                  #
#   Author: Zankar Sanghavi                        #
#                                                  #
#   © Dot Hill Systems Corporation                 #
#                                                  #
####################################################

import os

###################################
#  Importing from other Directory
###################################
import sys
os.chdir('..')
c_path = os.getcwd()
sys.path.insert(0, r''+str(c_path)+'/Common Scripts')

import fixed_data_ICS
import report_functions
import extract_lists
import modify_word_docx
import user_inputs_ICS


###################################
#  Importing from Current Directory
###################################
sys.path.insert(0, r''+str(c_path)+'/IO Stress')
import pandas
import csv
import numpy as np
import log_functions
import extract_data
import time
import zipfile

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_TABLE_DIRECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.shared import Cm
import re


###################################
#  To call functions from different
#  files. 
###################################
lf=log_functions.Log_Functions
ed= extract_data.Extract_Data
ui= user_inputs_ICS.User_Inputs
fd= fixed_data_ICS.Fixed_Data
rf= report_functions.Report_Functions


class Generate_Final_Word:

    ####################################################
    #   This Function does:
    #   
    #   1.) Uses "unzip_pull_log" function to unzip and
    #   find the .logs file in the folder to read the 
    #   data.
    #
    #   2.) Checks Errors, extracts: Hardware, Software
    #   and PhyStats/DPSLD Informations. Generates two
    #   DPSLD files to write in WORD report. 
    #   
    #   3.) It writes required data, with formatting
    #   into a WORD report. 
    #   
    ####################################################
    def generate_final_report(file_name, file_name1, fixed_dir
                                , part_no, test_name, no_of_files,
                                csv_file_list, usr_before_list,
                                usr_after_list, fw_no_names,
                                chassis_names, cntrllr_names,
                                t_fw_type):
        
        document = Document(r''+str(fixed_dir)+'\\'+str(part_no)
        +str(test_name)+'.docx')
        
        document.add_page_break()
        f=no_of_files
        error_collection=[]
        for f in range(no_of_files):
        
            progress=(round((float(100/int(no_of_files))*f),2)) 
            # to show Report Progress
            
            print('\nReport Progess: ',progress,'%\n')
            
            ub_temp = str(no_of_files)+' - Before Log file'
            ub_data= lf.unzip_pull_log(usr_before_list[f], 'store', ub_temp) 
            #pull Before .logs file from zip folder
            
            ua_temp=str(no_of_files)+' - After Log file'
            ua_data= lf.unzip_pull_log(usr_after_list[f], 'store', ua_temp) 
            #pull After .logs file from zip folder
            
            
            ###################################
            #  Error check and data extraction
            ###################################
            [write_sum, read_sum, hw_list, host_list
                , sasmap_list]= ed.generate_data_tables(
                                csv_file_list[f], ub_data
                                , ua_data,fixed_dir)
            
            
            section = document.sections[-1]     # last section in document
            ###################
            #   Summary
            ###################
            #document.add_page_break()
            document.add_heading('72hrs IO stress test Summary'\
                     'for ' + str(fw_no_names[f]) + '\\'
                     + str(chassis_names[f]) +  '\\' 
                     + str(cntrllr_names[f]) + ' chassis',level=3)
                     

            temp_para=document.add_paragraph()
            
            paragraph_format = temp_para.paragraph_format
            paragraph_format.left_indent
            paragraph_format.left_indent = Inches(0.5)
            
            run = temp_para.add_run('Read error(s): '
                    +str(read_sum)+ '\nWrite error(s): '
                    +str(write_sum)+'\n')
                    
            font = run.font
            font.name = 'Courier New'
            font.size = Pt(11)

            
            ###################
            #   F2 Menu
            ###################
            document.add_heading('72hrs IO stress (F2 menu) in ' 
            + str(fw_no_names[f]) + '\\' + str(chassis_names[f]) 
            +  '\\'  + str(cntrllr_names[f]) + ' chassis' ,level=3)

            for i in range(len(hw_list)):
                temp_para= document.add_paragraph()
                temp_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = temp_para.add_run(hw_list[i])
                font = run.font
                font.name = 'Courier New'
                font.size = Pt(11)
                             
                     
                     
            ###################
            #   F3 Menu
            ###################
            document.add_page_break()
            
            document.add_heading('72hrs IO stress (F3 menu) in ' 
            + str(fw_no_names[f]) + '\\' + str(chassis_names[f]) 
            +  '\\'  + str(cntrllr_names[f]) + ' chassis',level=3)
             
            for i in range(len(sasmap_list)):
                temp_para1= document.add_paragraph() 
                run = temp_para1.add_run(sasmap_list[i])
                font = run.font
                font.name = 'Courier New'
                font.size = Pt(11)
                       

            ###########################
            #   PhyStats 
            #   1st Table
            #   Comparison Table
            ###########################
            document.add_page_break()
            
            document.add_heading('72hrs IO stress'+ 
                                '(PhyStats Comparison) in '
                + str(fw_no_names[f]) + '\\' 
                + str(chassis_names[f]) +  '\\'
                + str(cntrllr_names[f]) + ' chassis' ,level=3)

            file_data = pandas.read_csv(open(r''+ str(fixed_dir) 
                                        +str(file_name1) +'.csv')
                                        ,header=None)
            file_data = np.array(file_data)

            file_data = file_data[:,2:] # we don't want first 2 columns

            no_rows= len(file_data[:,0]) #finding number of rows
            no_columns= len(file_data[0,:]) #finding number of columns

            ###########################
            #   Creating a table and 
            #   setting its attributes.
            ###########################
            table = document.add_table(rows=no_rows+1
                ,cols=no_columns)
                
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.table_direction = WD_TABLE_DIRECTION.LTR
            table_style = document.styles["Normal"]
            table_font1=table_style.font
            table.style = 'Table Grid'
            table.autofit=True
            
            
            ###########################
            #   If not satisfied with
            #   auto fit use this!
            ###########################
            # rf.set_column_width(table.columns[0], Cm(1)) #SLOT
            # rf.set_column_width(table.columns[1], Cm(1.5)) #Vendor
            # rf.set_column_width(table.columns[2], Cm(3)) #Model No.
            # rf.set_column_width(table.columns[3], Cm(2)) #Serial No.
            # rf.set_column_width(table.columns[4], Cm(1.8)) #Capacity
            # rf.set_column_width(table.columns[5], Cm(1.2)) #FW Rev.

            # rf.set_column_width(table.columns[6], Cm(1.2)) #Invalid DWORD Count
            # rf.set_column_width(table.columns[7], Cm(1.2)) #Running Disparity Err Count
            # rf.set_column_width(table.columns[8], Cm(1.2)) #Loss of DWORD Sync
            # rf.set_column_width(table.columns[9], Cm(1.2)) #Phy Reset Problem

            # rf.set_column_width(table.columns[10], Cm(1.2)) #Invalid DWORD Count
            # rf.set_column_width(table.columns[11], Cm(1.2)) #Running Disparity Err Count
            # rf.set_column_width(table.columns[12], Cm(1.2)) #Loss of DWORD Sync
            # rf.set_column_width(table.columns[13], Cm(1.2)) #Phy Reset Problem

            hdr_cells=table.rows[0].cells

            hdr_cells[0].text = '72 hr I/O Stress Phy counts'
            hdr_cells[0].merge(hdr_cells[5])

            hdr_cells[6].text = 'A Port(0)'
            hdr_cells[6].merge(hdr_cells[9])

            hdr_cells[10].text = 'B Port(1)'
            hdr_cells[10].merge(hdr_cells[13])


            for i in range(1,no_rows+1):
                for j in range(no_columns):
                    table_font1.size = Pt(5)
                    #font.size = Pt(8)
                    hdr_cells=table.rows[i].cells
                    hdr_cells[j].text = file_data[i-1,j]    
     
            document.add_paragraph('\n')


            ##########################
            #   DPSLD 
            #   2nd Table
            #   Before & After File
            ##########################
            document.add_page_break()
            document.add_heading('72hrs IO stress'
                +' (Log Files, DPSLD data) in ' 
                + str(fw_no_names[f]) + '\\' 
                + str(chassis_names[f])
                +  '\\' + str(cntrllr_names[f]) 
                + ' chassis',level=3)
                   
            file_data = pandas.read_csv(open(r''+ str(fixed_dir) +str(file_name) +'.csv'),header=None)
            file_data = np.array(file_data)

            no_rows= len(file_data[:,0])
            no_columns= len(file_data[0,:])

            section.start_type = WD_SECTION.NEW_PAGE # For Wide section

            section.top_margin= Inches(0.3)
            section.bottom_margin=Inches(0)

            section.left_margin = Inches(0.4)
            section.right_margin = Inches(0.4)

            
            ###########################
            #   Creating a table and 
            #   setting its attributes.
            ###########################
            table = document.add_table(rows=no_rows+1
                                       ,cols=no_columns)
                                       
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.table_direction = WD_TABLE_DIRECTION.LTR
            table_style = document.styles["Normal"]
            table_font=table_style.font

            table.autofit=False
            rf.set_column_width(table.columns[0], Cm(1.2)) 
            #SLOT
            rf.set_column_width(table.columns[1], Cm(0.5)) 
            #PHY
            rf.set_column_width(table.columns[2], Cm(1.6)) 
            #Invalid DWORD Count
            rf.set_column_width(table.columns[3], Cm(1.6)) 
            #Running Disparity Err Count
            rf.set_column_width(table.columns[4], Cm(1.6)) 
            #Loss of DWORD Sync
            rf.set_column_width(table.columns[5], Cm(1.6)) 
            #Phy Reset Problem
            rf.set_column_width(table.columns[6], Cm(1.8)) 
            #DFE Values

            rf.set_column_width(table.columns[7], Cm(1.2)) 
            #SLOT
            rf.set_column_width(table.columns[8], Cm(0.5)) 
            #PHY
            rf.set_column_width(table.columns[9], Cm(1.6)) 
            #Invalid DWORD Count
            rf.set_column_width(table.columns[10], Cm(1.6)) 
            #Running Disparity Err Count
            rf.set_column_width(table.columns[11], Cm(1.6)) 
            #Loss of DWORD Sync
            rf.set_column_width(table.columns[12], Cm(1.6)) 
            #Phy Reset Problem
            rf.set_column_width(table.columns[13], Cm(1.8)) 
            #DFE Values

            table.style = 'Table Grid'

            hdr_cells=table.rows[0].cells

            hdr_cells[0].text = 'BEFORE'
            hdr_cells[0].merge(hdr_cells[6])

            hdr_cells[7].text = 'AFTER'
            hdr_cells[7].merge(hdr_cells[13])

            for i in range(1,no_rows+1):
                for j in range(no_columns):
                    table_font.size=Pt(9)
                    hdr_cells=table.rows[i].cells
                    hdr_cells[j].text = file_data[i-1,j]  
            
            if f<no_of_files-1:
                document.add_page_break()
                
        document.save(r''+str(fixed_dir)+'\\'+str(part_no)
                      +str(test_name)+'.docx')
        return error_collection
        
       
#####################################
#              END                  #
#####################################